"""生成《电商双边智能预算规划系统 — 使用指南》DOCX 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ========== 全局样式 ==========
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_code_block(doc, text):
    """添加灰色背景的代码块段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # 灰色底纹
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F0F0')
    shading.set(qn('w:val'), 'clear')
    p.runs[0].element.rPr.append(shading)


def add_bullet(doc, text, bold_prefix=None):
    """添加带可选加粗前缀的要点"""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_warning_box(doc, title, text):
    """添加警告样式的段落"""
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️ {title}：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    p.add_run(text)


# ================================================================
# 封面
# ================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('电商双边智能预算规划系统', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('使用指南 & 部署手册')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(18)
subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph(f'版本：v1.0\n日期：{datetime.date.today().strftime("%Y年%m月%d日")}\n适用环境：Windows + VS Code + Python 3.12')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ================================================================
# 目录概要
# ================================================================
doc.add_heading('目录', level=1)
toc_items = [
    '一、环境准备与启动 — 拿到代码后怎么跑起来',
    '二、常见问题与避坑指南 — 基于实际踩坑经验',
    '三、网页功能详解 — 每个区域代表什么、怎么用',
    '    3.1  左侧栏：全局调度配置',
    '    3.2  左侧主区：购物车数据管理',
    '    3.3  右侧主区：算法决策大盘',
    '    3.4  底部：算法性能曲线分析',
    '四、典型使用流程',
    '附录：完整代码结构速览',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ================================================================
# 一、环境准备与启动
# ================================================================
doc.add_heading('一、环境准备与启动', level=1)
doc.add_heading('1.1 前置条件', level=2)
add_bullet(doc, 'Python 3.9+ 已安装（验证：终端输入 ', bold_prefix='')
add_code_block(doc, 'python --version')
add_bullet(doc, 'VS Code 或任意代码编辑器已安装', bold_prefix='')
add_bullet(doc, '代码文件 app.py 已拷贝到本地某个文件夹（比如 E:\\Algorithm）', bold_prefix='')

doc.add_heading('1.2 安装依赖', level=2)
doc.add_paragraph('打开 VS Code 终端（快捷键 Ctrl+`），执行以下命令安装所需的 Python 包：')
add_code_block(doc, 'pip install streamlit pandas matplotlib')

doc.add_paragraph('如果安装较慢，可使用国内镜像：')
add_code_block(doc, 'pip install streamlit pandas matplotlib -i https://pypi.tuna.tsinghua.edu.cn/simple')

doc.add_heading('1.3 启动网站', level=2)
doc.add_paragraph('在终端中，先 cd 到 app.py 所在目录，再启动：')
add_code_block(doc, 'cd E:\\Algorithm')
add_code_block(doc, 'python -m streamlit run app.py')

doc.add_paragraph()
doc.add_paragraph('启动成功后终端会显示：')
add_code_block(doc, '  You can now view your Streamlit app in your browser.')
add_code_block(doc, '  Local URL: http://localhost:8501')

doc.add_paragraph('此时在浏览器打开 http://127.0.0.1:8501 即可看到系统界面。')
doc.add_paragraph('按 Ctrl+C 可停止服务。')

doc.add_page_break()

# ================================================================
# 二、常见问题与避坑指南
# ================================================================
doc.add_heading('二、常见问题与避坑指南', level=1)
doc.add_paragraph('以下是本人在实际部署过程中遇到并解决的所有问题，请逐一阅读，可节省大量排查时间。')

doc.add_heading('2.1 问题①：终端输入 streamlit run app.py 报错"无法识别"', level=2)

doc.add_paragraph('错误信息：')
add_code_block(doc, "streamlit : 无法将\"streamlit\"项识别为 cmdlet、函数、脚本文件或可运行程序的名称。")

doc.add_paragraph('原因：')
doc.add_paragraph('Windows 下 Python 的 Scripts 目录未加入系统 PATH 环境变量，导致 streamlit.exe 无法被终端直接找到。')

doc.add_paragraph('解决方案（二选一）：')
add_bullet(doc, '方法 A（推荐）：每次启动都用 python -m 方式，无需改系统配置', bold_prefix='')
add_code_block(doc, 'python -m streamlit run app.py')
add_bullet(doc, '方法 B（一劳永逸）：将 Python Scripts 目录加入 PATH，之后可直接用 streamlit 命令', bold_prefix='')
add_code_block(doc, '[Environment]::SetEnvironmentVariable("Path", $env:Path + ";C:\\Users\\你的用户名\\AppData\\Local\\Microsoft\\WindowsApps\\Scripts", [EnvironmentVariableTarget]::User)')
doc.add_paragraph('执行后关闭并重新打开终端即可生效。')

doc.add_heading('2.2 问题②：网页打开后报错"No such file or directory ... app.zip.a01\\app.py"', level=2)

doc.add_paragraph('错误信息示例：')
add_code_block(doc, "[Errno 2] No such file or directory: 'C:\\\\Users\\\\...\\\\Temp\\\\bf37fdda-...app.zip.a01\\\\app.py'")

doc.add_paragraph('原因：')
doc.add_paragraph('端口 8501 上有之前残留的 Streamlit 进程在运行，那个旧进程是从临时解压目录启动的（比如你之前直接双击 zip 里的 app.py 运行过），文件路径指向已不存在的临时文件夹。你的浏览器连接到的是旧进程，而非你新启动的进程。')

doc.add_paragraph('解决方案：')
add_bullet(doc, '① 强制结束所有 Python 进程：', bold_prefix='')
add_code_block(doc, 'taskkill /F /IM python.exe')
add_bullet(doc, '② 确认端口已释放：', bold_prefix='')
add_code_block(doc, 'netstat -ano | findstr ":8501"')
doc.add_paragraph('  如果仍有 LISTENING 状态的条目，记下最后一列的 PID，然后 taskkill /F /PID xxxx 逐个杀掉。')
add_bullet(doc, '③ 重新启动：', bold_prefix='')
add_code_block(doc, 'python -m streamlit run app.py')
add_bullet(doc, '④ 打开浏览器，用 无痕模式（Ctrl+Shift+N）访问 http://127.0.0.1:8501', bold_prefix='')
doc.add_paragraph('  （无痕模式可避免浏览器缓存连接到旧进程）')

doc.add_heading('2.3 问题③：浏览器刷新后仍然显示旧错误', level=2)
doc.add_paragraph('原因：浏览器缓存了旧的错误页面，或 WebSocket 连接维持到了旧进程。')
doc.add_paragraph('解决方法：')
add_bullet(doc, '硬刷新：Ctrl + Shift + R（或 Ctrl + F5）')
add_bullet(doc, '或用无痕/隐私模式打开：Ctrl + Shift + N')
add_bullet(doc, '如果还不行，关闭所有同标签页 → 执行 taskkill → 重新启动 → 再打开')

doc.add_heading('2.4 问题④：图表中文显示为方框（乱码）', level=2)
doc.add_paragraph('原因：matplotlib 默认字体不支持中文。')
doc.add_paragraph('本代码已内置修复（app.py 第 8-10 行），会自动设置为 Microsoft YaHei 字体。如果仍有问题，确认系统安装了中文字体。')

doc.add_heading('2.5 问题⑤：性能测试时页面卡死或报内存不足', level=2)
doc.add_paragraph('原因：DP 算法复杂度为 O(N × W × B)，其中 W 是用户预算、B 是平台补贴。如果商品数量 × 预算太大，计算量和内存会爆炸。')
doc.add_paragraph('预防措施：')
add_bullet(doc, '性能测试的商品起点和终点从 小 → 大 逐步尝试（建议先 10 → 50）')
add_bullet(doc, '左侧栏预算不要一开始就拉到最大值')
add_bullet(doc, '如果卡死：Ctrl+C 终止 → taskkill 清进程 → 减小参数 → 重新启动')

doc.add_page_break()

# ================================================================
# 三、网页功能详解
# ================================================================
doc.add_heading('三、网页功能详解', level=1)

doc.add_paragraph('打开 http://127.0.0.1:8501 后，网页被分为以下区域。下面按 "从上到下、从左到右" 的顺序逐一说明。')

# ---- 3.1 左侧栏 ----
doc.add_heading('3.1 左侧栏 — 系统全局调度配置', level=2)
doc.add_paragraph('这是整个系统的"总控开关"，左侧栏的参数会同时影响右侧推荐结果和底部性能测试。')

doc.add_heading('① 用户消费预算（默认 1000 元）', level=3)
add_bullet(doc, '模拟双十一用户手里有多少满减券/红包可以花')
add_bullet(doc, '预算越高 → 算法能买的商品越多 → 推荐结果越丰富')
add_bullet(doc, '预算越低 → 只能挑最划算的几件')

doc.add_heading('② 平台补贴资金池（默认 150 元）', level=3)
add_bullet(doc, '模拟平台（淘宝/京东）愿意为这次促销掏出多少补贴帮用户抵钱')
add_bullet(doc, '范围 50 ~ 1000 元，每个商品都有一项"平台补贴"成本，选中越多消耗越多')
add_bullet(doc, '用户预算 + 平台补贴，两个约束同时起作用')

doc.add_heading('③ 多目标业务权重调节（默认 0.5 : 0.5）', level=3)
doc.add_paragraph('这是系统的灵魂参数，决定算法到底"偏向谁"。公式为：')
doc.add_paragraph('    复合价值 = α × 用户喜爱度 + β × 商家佣金', style='Intense Quote')
add_bullet(doc, 'α 越大（滑块往右拉）→ 算法更看重"用户喜不喜欢"，推荐用户最爱但利润可能低的商品')
add_bullet(doc, 'α 越小（滑块往左拉）→ 算法更看重"商家赚多少佣金"，推荐利润高但用户未必喜欢的商品')
add_bullet(doc, 'β 会自动计算，始终 = 1 - α，右侧会实时显示')

# ---- 3.2 左侧主区 ----
doc.add_heading('3.2 左侧主区 — 购物车数据管理', level=2)
doc.add_paragraph('这里是算法的"原材料"——候选商品池。算法会从中挑选出一部分推荐。')

doc.add_heading('📦 商品表格', level=3)
add_bullet(doc, '展示所有候选商品，默认 6 件')
add_bullet(doc, '共 6 列：商品名称 | 商家 | 用户原价 | 平台补贴 | 用户喜爱度 | 商家佣金')
add_bullet(doc, '每一行是一个候选商品，算法将从这些候选中挑选')

doc.add_heading('➕ 手动添加商品（可折叠面板）', level=3)
doc.add_paragraph('点开折叠面板，填写以下字段后点击「添加到购物车」：')
add_bullet(doc, '商品名称：随便填，如"无线鼠标"')
add_bullet(doc, '所属商家：品牌/店铺名，如"雷蛇旗舰店"')
add_bullet(doc, '用户原价：用户要付多少钱')
add_bullet(doc, '平台补贴：平台帮你贴多少')
add_bullet(doc, '用户喜爱度：1-100 的整数，越高用户越想要')
add_bullet(doc, '商家佣金：平台从这笔订单能赚多少')
doc.add_paragraph('添加后页面自动刷新，表格会多出一行。')

doc.add_heading('✨ 批量生成随机商品（可折叠面板）', level=3)
doc.add_paragraph('不想手动填？输入数量（比如 20），点按钮，系统会用随机数据 替换 当前所有商品。适用于快速压测算法性能。')

# ---- 3.3 右侧主区 ----
doc.add_heading('3.3 右侧主区 — 算法决策大盘', level=2)
doc.add_paragraph('这是核心结果区。每次你调整左侧参数或购物车数据，这里会自动重新计算。')

doc.add_heading('📊 两种算法得分对比（三列卡片）', level=3)
doc.add_paragraph('右侧首先展示三列指标卡片，帮你快速判断两种算法的差距：')
add_bullet(doc, '🏆 DP 动态规划（数学最优解）：DP 算出的得分，理论上不可能更高', bold_prefix='')
add_bullet(doc, '⚡ 贪心算法（近似解）：贪心算出的得分，速度快但不保证最优', bold_prefix='')
add_bullet(doc, '📉 贪心落后 DP 多少：两者的差距，0% = 平手，数值越大贪心越差', bold_prefix='')
doc.add_paragraph('卡片下方还有一条彩色结论条，自动判断贪心在这个场景下是否够用：')
add_bullet(doc, '差距 = 0%：绿色 — 贪心也找到了最优解')
add_bullet(doc, '差距 < 5%：绿色 — 贪心质量很高，速度优势明显')
add_bullet(doc, '差距 ≥ 5%：黄色 — 追求收益时建议用 DP，追求速度时贪心可接受')

doc.add_heading('🤖 方案 A：DP 动态规划（Tab 页）', level=3)
doc.add_paragraph('展示 DP 算法的详细结果：')
add_bullet(doc, '4 个指标：计算耗时（秒）、总复合价值得分、消费者总满意度、平台总佣金收益')
add_bullet(doc, '2 条进度条：用户预算占用（如 800/1000）、平台补贴占用（如 120/150）')
add_bullet(doc, '一张表格：DP 具体推荐了哪些商品')

doc.add_heading('⚡ 方案 B：贪心算法（Tab 页）', level=3)
doc.add_paragraph('布局和方案 A 一模一样，但展示的是贪心算法的结果。切换两个 Tab 对比，就能直观看到 DP 和贪心分别选了哪些商品、哪边得分更高。')

# ---- 3.4 底部 ----
doc.add_heading('3.4 底部 — 算法性能曲线分析', level=2)
doc.add_paragraph('这是独立于上方推荐结果的性能实验区。')

doc.add_heading('和上方的关系', level=3)
add_bullet(doc, '沿用左侧栏的预算和权重设置')
add_bullet(doc, '但不使用上方购物车的商品 — 而是按指定数量随机生成新商品来测试')
add_bullet(doc, '目的：看两种算法随着商品变多，耗时分别怎么增长')

doc.add_heading('三个滑块（控制 x 轴取哪些测试点）', level=3)
add_bullet(doc, '从多少件商品开始测：x 轴起点，默认 10', bold_prefix='')
add_bullet(doc, '最多测到多少件商品：x 轴终点，默认 50', bold_prefix='')
add_bullet(doc, '每隔多少件商品测一次：x 轴步长，默认 10', bold_prefix='')
doc.add_paragraph()
doc.add_paragraph('示例（10→50，步长 10）：系统用 10、20、30、40、50 这 5 个数据点分别跑两种算法，记录每次耗时。')

add_warning_box(doc, '注意', 'DP 复杂度为 O(N×W×B)，商品数量或预算太大时计算可能非常慢。建议从 10→50 开始试，确认不卡再增大。')

doc.add_heading('点击「运行基准测试」后', level=3)
add_bullet(doc, '出现进度条，显示当前测试进度')
add_bullet(doc, '完成后展示数据表格：每个测试点的精确耗时 + "DP 比贪心慢多少倍"')
add_bullet(doc, '以及一张对数坐标折线图，每个数据点旁边标注精确秒数')

doc.add_heading('怎么看结果', level=3)
add_bullet(doc, '蓝线（DP）：商品越多越陡峭，计算量增长很快')
add_bullet(doc, '红线（贪心）：几乎平躺，因为排序后扫一遍即可，非常快')
doc.add_paragraph('这就是算法设计中的经典权衡：DP 追求最优但慢，贪心追求速度但有质量损失。表格中的"DP/贪心 倍数"直观量化了这种差距。')

doc.add_page_break()

# ================================================================
# 四、典型使用流程
# ================================================================
doc.add_heading('四、典型使用流程', level=1)

doc.add_paragraph('推荐按以下步骤操作：')

steps = [
    ('① 调整左侧栏',
     '设定预算（用户 1000、平台 150）→ 调节权重（α=0.5 均衡；α=0.7 偏向用户；α=0.3 偏向平台）'),
    ('② 准备商品数据',
     '用默认 6 件商品看效果，或用「批量生成」功能生成 10~20 件随机商品'),
    ('③ 查看推荐结果',
     '对比三列卡片中 DP 和贪心的得分差距 → 切换两个 Tab 看各自选了哪些商品 → 看彩色结论条判断贪心是否够用'),
    ('④ 跑性能测试（可选）',
     '设置 10→50→步长 10 → 点击「运行基准测试」→ 等进度条跑完 → 看图 + 表格了解算法效率'),
    ('⑤ 换参数重试',
     '改权重或预算 → 右侧自动重新计算 → 观察得分变化 → 再次跑性能测试对比'),
]

for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(desc)

doc.add_page_break()

# ================================================================
# 附录
# ================================================================
doc.add_heading('附录：完整代码结构速览', level=1)
doc.add_paragraph('app.py 共约 400 行，分为以下模块：')

modules = [
    ('第 1-10 行', '导入 & matplotlib 中文字体修复'),
    ('第 12-107 行', '核心算法引擎 — DualSidedECommerceSystem 类\n'
     '  · run_joint_optimization() — 二维 DP 背包，保证全局最优\n'
     '  · run_fallback_greedy()   — 贪心算法，按性价比排序取商品'),
    ('第 109-129 行', '辅助函数 — generate_random_items() 批量生成随机商品'),
    ('第 131-160 行', 'Streamlit 页面配置 & 侧边栏控件'),
    ('第 162-190 行', '左侧主区 — 购物车表格 & 商品添加/生成'),
    ('第 192-285 行', '右侧主区 — 算法得分对比 & DP/贪心 Tab 分页结果'),
    ('第 287-395 行', '底部 — 性能基准测试 & 耗时曲线图'),
]

for lines, desc in modules:
    p = doc.add_paragraph()
    run = p.add_run(f'{lines}：')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph()
final = doc.add_paragraph('— 文档结束 —')
final.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ========== 保存 ==========
output_path = r'e:\Algorithm\使用指南_电商双边智能预算规划系统.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
